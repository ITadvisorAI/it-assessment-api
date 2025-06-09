
from docx import Document
from docx.shared import Inches
import os

def generate_docx_report(session_id, hw_df, sw_df, chart_paths):
    """Generate a detailed DOCX report.

    Parameters
    ----------
    session_id : str
        Unique identifier for the current assessment session.
    hw_df : pandas.DataFrame
        Hardware dataframe to summarize.
    sw_df : pandas.DataFrame
        Software dataframe to summarize.
    chart_paths : dict
        Dictionary of chart names to file paths returned by
        :func:`visualization.generate_charts`.

    Returns
    -------
    str or None
        Path to the generated DOCX file or ``None`` if generation failed.
    """
    try:
        output_path = os.path.join("temp_sessions", session_id, "IT_Current_Status_Assessment_Report.docx")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        document = Document()
        document.add_heading('IT Infrastructure Current Status Report', 0)

        document.add_heading('Session ID', level=1)
        document.add_paragraph(session_id)

        document.add_heading('Hardware Summary', level=1)
        if hw_df is not None and not hw_df.empty:
            table = document.add_table(rows=1, cols=len(hw_df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(hw_df.columns):
                hdr_cells[i].text = str(col)
            for _, row in hw_df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
        else:
            document.add_paragraph("No hardware data available.")

        document.add_heading('Software Summary', level=1)
        if sw_df is not None and not sw_df.empty:
            table = document.add_table(rows=1, cols=len(sw_df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(sw_df.columns):
                hdr_cells[i].text = str(col)
            for _, row in sw_df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
        else:
            document.add_paragraph("No software data available.")

        document.add_heading('Charts & Visualizations', level=1)
        for chart_path in chart_paths.values():
            if os.path.exists(chart_path):
                document.add_picture(chart_path, width=Inches(5.5))
            else:
                document.add_paragraph(f"⚠️ Missing chart: {chart_path}")

        document.save(output_path)
        return output_path

    except Exception as e:
        print("❌ Error in generate_docx_report:", str(e))
        return None
